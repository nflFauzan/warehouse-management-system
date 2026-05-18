
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Page margins ──────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(3)
    section.bottom_margin = Cm(3)
    section.left_margin   = Cm(4)
    section.right_margin  = Cm(3)

IMG_DIR = os.path.join(os.path.dirname(__file__), "laporan_images")

def add_image(filename, width=Inches(5.5)):
    path = os.path.join(IMG_DIR, filename)
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=width)
    else:
        p = doc.add_paragraph(f"[Gambar: {filename} — letakkan file di docs/laporan_images/]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].italic = True
        p.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)

def heading1(text):
    h = doc.add_heading(text, level=1)
    h.runs[0].font.color.rgb = RGBColor(0x1F, 0x39, 0x7A)

def heading2(text):
    h = doc.add_heading(text, level=2)
    h.runs[0].font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

def heading3(text):
    doc.add_heading(text, level=3)

def para(text, bold=False, italic=False, size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    return p

def add_table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        tc = hdr[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "2E75B6")
        tcPr.append(shd)
        for run in hdr[i].paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for ri, row in enumerate(rows):
        cells = t.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            cells[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    return t

# ══════════════════════════════════════════════════════════════
#  COVER
# ══════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("LAPORAN USABILITY TESTING")
r.bold = True; r.font.size = Pt(18); r.font.name = "Times New Roman"
r.font.color.rgb = RGBColor(0x1F, 0x39, 0x7A)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run("Warehouse Management System (WMS) TAKKA STEEL")
r2.bold = True; r2.font.size = Pt(14); r2.font.name = "Times New Roman"

doc.add_paragraph()

for line in [
    "Mata Kuliah: Interaksi Manusia dan Komputer",
    "Program Studi: Sistem Informasi",
    "",
    "Disusun oleh:",
    "1. Muhamad Naufal Fauzan   — 241572010008",
    "2. Siti Tahtia Ainun Zahra — 241572010014",
    "3. Mutiara Adinda          — 241572010005",
    "4. Rahma Fitria Tunnisa    — 241572010009",
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(line)
    r.font.size = Pt(12); r.font.name = "Times New Roman"
    if line.startswith("Disusun"):
        r.bold = True

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("2024 / 2025")
r.bold = True; r.font.size = Pt(12); r.font.name = "Times New Roman"

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#  BAB 1 — PENDAHULUAN
# ══════════════════════════════════════════════════════════════
heading1("BAB 1 — PENDAHULUAN")

heading2("1.1 Latar Belakang")
para(
    "Sistem informasi berbasis web semakin banyak digunakan dalam pengelolaan operasional "
    "bisnis, termasuk manajemen pergudangan. WMS TAKKA STEEL merupakan sistem yang dirancang "
    "untuk membantu staff dan pemilik toko baja TAKKA STEEL dalam mengelola data barang, "
    "transaksi masuk-keluar, posisi stok, serta laporan secara digital dan terpusat. "
    "Sebelum sistem dikembangkan secara penuh, dilakukan serangkaian tahap perancangan "
    "antarmuka yang meliputi pembuatan Information Architecture (IA), Wireframe, dan Prototype "
    "interaktif menggunakan Figma. Sebagai tahap validasi desain, dilakukan Usability Testing "
    "terhadap prototype tersebut untuk memastikan kemudahan navigasi dan ketepatan alur "
    "penggunaan sebelum masuk ke tahap pengembangan lebih lanjut."
)

heading2("1.2 Tujuan")
para("Tujuan dari kegiatan usability testing ini adalah:")
for t in [
    "Mengevaluasi kemudahan navigasi antar halaman pada prototype WMS TAKKA STEEL.",
    "Mengidentifikasi kendala atau kebingungan pengguna dalam menemukan fitur yang dibutuhkan.",
    "Memperoleh insight untuk perbaikan desain antarmuka sebelum tahap implementasi.",
    "Memastikan alur penggunaan sistem sesuai dengan kebutuhan dua jenis pengguna utama: Staff Gudang dan Owner.",
]:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(t).font.size = Pt(12)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#  BAB 2 — PERANCANGAN SISTEM
# ══════════════════════════════════════════════════════════════
heading1("BAB 2 — PERANCANGAN SISTEM")

# 2.1 IA
heading2("2.1 Information Architecture (IA)")
para(
    "Information Architecture (IA) adalah struktur hierarki yang menggambarkan organisasi "
    "konten dan navigasi dalam sebuah sistem. IA WMS TAKKA STEEL dirancang untuk mencakup "
    "seluruh modul operasional gudang dengan pemisahan yang jelas antara modul Master Data, "
    "Transaksi, Stok, Warehouse Layout, Laporan, dan Settings."
)
para(
    "Sistem memiliki dua role pengguna utama: Staff Gudang yang dapat mengakses modul "
    "operasional (transaksi, stok, warehouse layout), dan Owner yang memiliki akses penuh "
    "termasuk laporan dan manajemen pengguna. Berikut adalah diagram IA WMS TAKKA STEEL:"
)
doc.add_paragraph()
add_image("ia_diagram.png", width=Inches(6.0))
p = doc.add_paragraph("Gambar 2.1 — Information Architecture WMS TAKKA STEEL")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].italic = True; p.runs[0].font.size = Pt(10)

doc.add_paragraph()
para("Struktur IA terdiri dari modul-modul berikut:")
ia_rows = [
    ("Login", "Autentikasi pengguna dengan role Staff / Owner"),
    ("Dashboard", "Ringkasan KPI, grafik transaksi, dan inventaris terkini"),
    ("Master Data", "Pengelolaan Items, Kategori, Satuan, Supplier, Customer"),
    ("Transaksi", "Stock In (List, Form, Detail) dan Stock Out (List, Form, Detail)"),
    ("Stok", "Stock Position (posisi saat ini) dan Stock History (riwayat)"),
    ("Warehouse Layout", "Visualisasi peta rak gudang dan alokasi barang"),
    ("Laporan", "Stock Report, Stock In Report, Stock Out Report"),
    ("Settings", "Profile pengguna dan User Management (Owner only)"),
]
add_table(["Modul", "Deskripsi"], ia_rows)

doc.add_page_break()

# 2.2 Wireframe
heading2("2.2 Wireframe")
para(
    "Wireframe adalah representasi visual low-fidelity yang menggambarkan tata letak dan "
    "struktur antarmuka tanpa elemen visual detail seperti warna dan gambar. Wireframe "
    "WMS TAKKA STEEL dibuat dalam format grayscale untuk memfokuskan evaluasi pada "
    "hierarki informasi, penempatan komponen, dan alur navigasi."
)
para(
    "Tiga halaman utama dipilih sebagai wireframe representatif karena mencakup "
    "interaksi terpenting dalam sistem:"
)

heading3("2.2.1 Dashboard")
para(
    "Halaman Dashboard dirancang sebagai pusat informasi utama yang menampilkan 4 kartu KPI "
    "(Key Performance Indicator) di bagian atas, area grafik transaksi di tengah, panel "
    "Inventory Overview di sisi kanan, serta tabel Recent Transactions di bagian bawah. "
    "Navigasi sidebar di sisi kiri memuat seluruh menu utama sistem."
)
doc.add_paragraph()
add_image("wireframe_dashboard.png")
p = doc.add_paragraph("Gambar 2.2 — Wireframe Halaman Dashboard")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].italic = True; p.runs[0].font.size = Pt(10)

doc.add_paragraph()
heading3("2.2.2 Stock In Form (Formulir Barang Masuk)")
para(
    "Halaman Stock In Form menampilkan formulir pencatatan transaksi barang masuk dengan "
    "field: Reference No (auto-generated), Supplier, Tanggal, Warehouse, dan Catatan. "
    "Bagian Item Details memuat tabel barang yang bisa ditambah secara dinamis dengan "
    "informasi Nama Barang, Satuan, Qty, Harga Satuan, dan Subtotal. Di bagian bawah "
    "ditampilkan ringkasan Grand Total beserta tombol Cancel dan Save Transaction."
)
doc.add_paragraph()
add_image("wireframe_stockin.png")
p = doc.add_paragraph("Gambar 2.3 — Wireframe Halaman Stock In Form")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].italic = True; p.runs[0].font.size = Pt(10)

doc.add_paragraph()
heading3("2.2.3 Warehouse Layout (Peta Gudang)")
para(
    "Halaman Warehouse Layout menyajikan visualisasi peta rak gudang dalam format grid "
    "dengan kode rak (A-01 s.d. D-05). Setiap rak menampilkan nama item yang tersimpan "
    "dan persentase kapasitas terpakai dengan indikator visual (kosong/partial/penuh/nonaktif). "
    "Panel kanan menampilkan daftar Unallocated Items yang belum dialokasikan ke rak, "
    "dilengkapi tombol Allocate dan Auto-Allocate All."
)
doc.add_paragraph()
add_image("wireframe_warehouse.png")
p = doc.add_paragraph("Gambar 2.4 — Wireframe Halaman Warehouse Layout")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].italic = True; p.runs[0].font.size = Pt(10)

doc.add_page_break()

# 2.3 Prototype
heading2("2.3 Prototype")
para(
    "Prototype interaktif dibuat menggunakan Figma sebagai alat desain berbasis web. "
    "Prototype ini bersifat read-only navigable — pengguna dapat berpindah antar halaman "
    "melalui klik elemen yang telah dihubungkan, namun belum mendukung input data secara "
    "fungsional. Tujuannya adalah menguji navigabilitas dan kejelasan alur antar halaman "
    "sebelum implementasi sistem yang sesungguhnya."
)
para(
    "Link Prototype Figma: https://www.figma.com/design/goiqc24dXo66F8im6eH9Cf/Untitled"
    "?node-id=22-2434&t=cpGE9tNVhxpyzuvS-0",
    italic=True
)
para(
    "Prototype mencakup seluruh halaman yang terdefinisi dalam IA, dengan koneksi navigasi "
    "antar halaman mengikuti alur penggunaan yang realistis. Keterbatasan prototype saat ini "
    "adalah belum dapat melakukan input/submit data; oleh karena itu, usability testing "
    "difokuskan pada evaluasi navigasi dan kejelasan informasi antarmuka."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#  BAB 3 — USABILITY TESTING
# ══════════════════════════════════════════════════════════════
heading1("BAB 3 — USABILITY TESTING")

heading2("3.1 Metode")
para(
    "Metode yang digunakan adalah Moderated Usability Testing dengan teknik Think-Aloud. "
    "Setiap partisipan diminta untuk menyelesaikan task scenario sambil mengungkapkan "
    "pikirannya secara lisan. Moderator membacakan instruksi tanpa memberikan petunjuk "
    "arah, sementara Observer mencatat kendala dan waktu penyelesaian."
)
para(
    "Karena prototype bersifat navigasi-only, pengujian difokuskan pada aspek "
    "Navigation & Findability: kemampuan pengguna menemukan halaman yang tepat "
    "untuk menyelesaikan tugas yang diberikan."
)
add_table(
    ["Aspek", "Detail"],
    [
        ("Jenis Testing", "Moderated Usability Testing"),
        ("Teknik", "Think-Aloud Protocol"),
        ("Fokus Evaluasi", "Navigasi, Findability, Kejelasan Label Menu"),
        ("Media Prototype", "Figma Interactive Prototype (read-only navigable)"),
        ("Durasi per sesi", "± 20–30 menit"),
        ("Jumlah Task", "5 task scenario"),
    ]
)

doc.add_paragraph()
heading2("3.2 Partisipan")
para(
    "Testing dilakukan oleh 4 anggota kelompok dengan pembagian peran sebagai berikut:"
)
add_table(
    ["No", "Nama", "NIM", "Peran dalam Testing", "Role Sistem"],
    [
        ("1", "Muhamad Naufal Fauzan",   "241572010008", "Moderator", "—"),
        ("2", "Siti Tahtia Ainun Zahra", "241572010014", "Observer",  "—"),
        ("3", "Mutiara Adinda",          "241572010005", "Partisipan / User 1", "Staff Gudang"),
        ("4", "Rahma Fitria Tunnisa",    "241572010009", "Partisipan / User 2", "Owner"),
    ]
)
doc.add_paragraph()
para(
    "Partisipan dipilih berdasarkan kedekatan profil dengan pengguna target sistem: "
    "User 1 merepresentasikan petugas operasional gudang (Staff), sedangkan User 2 "
    "merepresentasikan pemilik atau pengelola bisnis (Owner) yang membutuhkan akses "
    "ke laporan dan manajemen pengguna."
)

doc.add_paragraph()
heading2("3.3 Task Scenario")
para(
    "Lima task scenario dirancang untuk menguji navigasi pada fitur-fitur inti WMS. "
    "Setiap task diformulasikan dalam konteks peran pengguna yang realistis:"
)
add_table(
    ["Task", "Instruksi kepada Partisipan", "Halaman Target", "Role"],
    [
        ("T1", "Kamu baru masuk ke sistem. Di mana kamu akan mulai jika ingin melihat ringkasan kondisi gudang hari ini?",
         "Dashboard", "Staff / Owner"),
        ("T2", "Ada barang baja baru datang dari supplier. Navigasikan ke halaman yang kamu gunakan untuk mencatat barang masuk tersebut.",
         "Transactions → Stock In", "Staff"),
        ("T3", "Kamu ingin tahu berapa stok baja yang tersisa saat ini dan di rak mana posisinya. Temukan halamannya.",
         "Stock → Stock Position", "Staff / Owner"),
        ("T4", "Sebagai owner, kamu ingin melihat laporan seluruh stok barang bulan ini. Navigasikan ke halaman yang tepat.",
         "Reports → Stock Report", "Owner"),
        ("T5", "Kamu ingin melihat peta visual rak gudang dan mengecek rak mana yang sudah penuh. Temukan halamannya.",
         "Warehouse Layout", "Staff / Owner"),
    ]
)

doc.add_page_break()
heading2("3.4 Hasil Pengujian")
para(
    "Berikut adalah hasil observasi untuk setiap task, mencatat apakah partisipan "
    "berhasil menemukan halaman target, jalur navigasi yang ditempuh, dan kendala "
    "yang ditemui:"
)

heading3("Task 1 — Ringkasan Kondisi Gudang (Dashboard)")
add_table(
    ["Partisipan", "Status", "Waktu", "Catatan"],
    [
        ("Mutiara (User 1)", "Berhasil", "~12 dtk", "Langsung klik Dashboard di sidebar tanpa hambatan"),
        ("Rahma (User 2)",   "Berhasil", "~15 dtk", "Langsung klik Dashboard di sidebar tanpa hambatan"),
    ]
)
doc.add_paragraph()

heading3("Task 2 — Navigasi ke Stock In Form (Barang Masuk)")
add_table(
    ["Partisipan", "Status", "Waktu", "Catatan"],
    [
        ("Mutiara (User 1)", "Berhasil",              "~20 dtk", "Langsung menuju Transactions, lalu klik Barang Masuk"),
        ("Rahma (User 2)",   "Berhasil (1x salah klik)", "~45 dtk",
         "Sempat klik menu 'Barang' di Master Data karena mengira 'Barang' adalah "
         "menu untuk mencatat barang masuk. Setelah tidak menemukan form transaksi, "
         "backtrack ke sidebar dan menemukan 'Barang Masuk' di menu Transactions."),
    ]
)
doc.add_paragraph()

heading3("Task 3 — Posisi Stok Barang")
add_table(
    ["Partisipan", "Status", "Waktu", "Catatan"],
    [
        ("Mutiara (User 1)", "Berhasil", "~18 dtk", "Langsung menuju menu Stock → Stock Position"),
        ("Rahma (User 2)",   "Berhasil", "~22 dtk", "Langsung menuju menu Stock → Stock Position"),
    ]
)
doc.add_paragraph()

heading3("Task 4 — Laporan Stok Bulanan")
add_table(
    ["Partisipan", "Status", "Waktu", "Catatan"],
    [
        ("Mutiara (User 1)", "Berhasil", "~16 dtk", "Navigasi langsung ke Reports → Stock Report"),
        ("Rahma (User 2)",   "Berhasil", "~20 dtk", "Navigasi langsung ke Reports → Stock Report"),
    ]
)
doc.add_paragraph()

heading3("Task 5 — Warehouse Layout (Peta Gudang)")
add_table(
    ["Partisipan", "Status", "Waktu", "Catatan"],
    [
        ("Mutiara (User 1)", "Berhasil", "~12 dtk", "Langsung klik menu Warehouse di sidebar"),
        ("Rahma (User 2)",   "Berhasil", "~14 dtk", "Langsung menemukan tanpa hambatan"),
    ]
)

doc.add_paragraph()
heading2("3.5 Ringkasan Hasil (Completion Rate)")
add_table(
    ["Task", "Deskripsi", "Mutiara", "Rahma", "Completion Rate", "Keterangan"],
    [
        ("T1", "Dashboard",        "Berhasil", "Berhasil", "100%", "Tidak ada hambatan"),
        ("T2", "Stock In Form",    "Berhasil", "Berhasil (1x salah)", "100%",
         "1 partisipan sempat salah klik menu 'Barang' (Master Data)"),
        ("T3", "Stock Position",   "Berhasil", "Berhasil", "100%", "Tidak ada hambatan"),
        ("T4", "Laporan Stok",     "Berhasil", "Berhasil", "100%", "Tidak ada hambatan"),
        ("T5", "Warehouse Layout", "Berhasil", "Berhasil", "100%", "Tidak ada hambatan"),
    ]
)
para(
    "Seluruh task berhasil diselesaikan oleh kedua partisipan dengan completion rate 100%. "
    "Satu-satunya kendala ditemukan pada Task 2, di mana 1 partisipan (User 2) sempat "
    "salah memilih menu 'Barang' di Master Data karena kesamaan kata dengan 'Barang Masuk' "
    "di menu Transactions. Setelah backtrack singkat, partisipan berhasil menemukan halaman yang tepat.",
    italic=True
)

doc.add_page_break()
heading2("3.6 Insight")
para(
    "Dari hasil pengujian, ditemukan 2 insight utama berdasarkan pola yang terobservasi "
    "selama sesi usability testing:"
)

insights = [
    ("Ambiguitas label 'Barang' vs 'Barang Masuk' menyebabkan salah navigasi pada Task 2",
     "Partisipan User 2 (Owner) sempat mengarahkan kursor ke menu 'Barang' di Master Data "
     "ketika diminta mencari halaman pencatatan barang masuk. Hal ini terjadi karena kata "
     "'Barang' yang muncul di menu Master Data terasa lebih dekat secara semantik dengan "
     "instruksi 'ada barang baru datang'. Pengguna tidak langsung mengasosiasikan tindakan "
     "'mencatat barang masuk' dengan menu 'Transactions'. Ini menunjukkan bahwa label menu "
     "Master Data (khususnya sub-item 'Barang') perlu dibedakan lebih jelas dari terminologi "
     "operasional (Barang Masuk / Barang Keluar) yang ada di menu Transactions."),
    ("Empat dari lima task diselesaikan tanpa hambatan — struktur navigasi sudah intuitif",
     "Task 1 (Dashboard), Task 3 (Stock Position), Task 4 (Laporan Stok), dan Task 5 "
     "(Warehouse Layout) berhasil diselesaikan kedua partisipan secara langsung tanpa "
     "salah klik. Hal ini mengonfirmasi bahwa penamaan dan pengelompokan menu di sidebar "
     "sudah cukup jelas dan sesuai dengan mental model pengguna untuk fitur-fitur tersebut. "
     "Secara khusus, menu Warehouse Layout diakui paling mudah ditemukan berkat kombinasi "
     "nama yang deskriptif dan ikon visual yang representatif."),
]
for i, (title, desc) in enumerate(insights, 1):
    p = doc.add_paragraph(style="List Number")
    r = p.add_run(f"{title}\n")
    r.bold = True; r.font.size = Pt(12)
    r2 = p.add_run(desc)
    r2.font.size = Pt(12)

doc.add_paragraph()
heading2("3.7 Rekomendasi Perbaikan Desain")
para(
    "Berdasarkan insight di atas, berikut adalah rekomendasi perbaikan desain yang "
    "disusun berdasarkan tingkat dampak dan urgensinya:"
)
add_table(
    ["No", "Rekomendasi", "Area", "Prioritas"],
    [
        ("1",
         "Bedakan label sub-menu Master Data dari terminologi operasional. "
         "Ubah 'Barang' menjadi 'Data Barang' atau 'Master Barang' agar tidak "
         "tertukar dengan 'Barang Masuk' di menu Transactions.",
         "Sidebar — Master Data", "Tinggi"),
        ("2",
         "Tampilkan sub-menu Transactions (Barang Masuk / Barang Keluar) secara "
         "langsung terlihat di sidebar tanpa perlu expand/hover terlebih dahulu, "
         "sehingga pengguna dapat melihat pilihan navigasi lebih cepat.",
         "Sidebar — Transactions", "Sedang"),
        ("3",
         "Pertahankan desain keseluruhan menu navigasi. Task 1, 3, 4, dan 5 "
         "diselesaikan tanpa hambatan, menunjukkan struktur navigasi sudah benar "
         "dan cukup intuitif untuk pengguna baru sekalipun.",
         "Global Navigation", "Pertahankan"),
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#  PENUTUP
# ══════════════════════════════════════════════════════════════
heading1("PENUTUP")
para(
    "Usability testing terhadap prototype WMS TAKKA STEEL telah berhasil dilaksanakan "
    "dengan melibatkan 2 partisipan yang merepresentasikan dua peran pengguna utama: "
    "Staff Gudang dan Owner. Seluruh task scenario berhasil diselesaikan dengan completion "
    "rate 100%, yang menunjukkan bahwa struktur navigasi dan tata letak antarmuka prototype "
    "secara umum sudah mendukung alur kerja pengguna dengan baik."
)
para(
    "Satu-satunya kendala yang ditemukan selama sesi testing terjadi pada Task 2, "
    "di mana satu partisipan sempat salah memilih menu 'Barang' di Master Data "
    "karena kesamaan kata dengan 'Barang Masuk' yang ada di menu Transactions. "
    "Kendala ini bersifat minor dan dapat diselesaikan dengan perbaikan label menu "
    "yang lebih deskriptif — khususnya dengan membedakan terminologi antara data master "
    "dan data operasional/transaksi."
)
para(
    "Dengan hasil ini, kelompok memperoleh validasi yang kuat bahwa desain antarmuka "
    "WMS TAKKA STEEL telah berada pada jalur yang tepat. Perbaikan yang direkomendasikan "
    "bersifat minimal dan terfokus, sehingga tidak memerlukan perubahan struktural besar "
    "pada rancangan yang sudah ada. Prototype ini siap untuk dikembangkan lebih lanjut "
    "ke tahap implementasi sistem dengan beberapa penyesuaian label navigasi yang telah "
    "diidentifikasi melalui sesi usability testing ini."
)

# Save
out = os.path.join(os.path.dirname(__file__), "Laporan_Usability_Testing_WMS_TAKKA_STEEL.docx")
doc.save(out)
print(f"[OK] Dokumen berhasil dibuat: {out}")
